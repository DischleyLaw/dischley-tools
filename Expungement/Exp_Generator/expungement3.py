import os
os.environ['TK_SILENCE_DEPRECATION'] = '1'
import tkinter as tk
from tkinter import filedialog, ttk
from docx import Document

# Predefined prosecutor information based on county
prosecutor_info = {
    "Fairfax County": {"name": "Steve Descano, Esq.", "title": "Fairfax County Commonwealth's Attorney", "address1": "4110 Chain Bridge Road", "address2": "Fairfax, VA 22030"},
    "Arlington County": {"name": "Parisa Tafti, Esq.", "title": "Arlington County Commonwealth's Attorney", "address1": "1425 N. Courthouse Rd", "address2": "Arlington, VA 22201"},
    "Prince William County": {"name": "Amy Ashworth, Esq.", "title": "Prince William County Commonwealth's Attorney", "address1": "9311 Lee Ave", "address2": "Manassas, VA 20110"},
    "Loudoun County": {"name": "Robert Anderson, Esq.", "title": "Loudoun County Commonwealth's Attorney", "address1": "20 E Market St", "address2": "Leesburg, VA 20176"},
    "City of Alexandria": {"name": "Bryan Porter, Esq.", "title": "Alexandria City Commonwealth's Attorney", "address1": "520 King Street", "address2": "Alexandria, VA 22314"},
    "Stafford County": {"name": "Eric Olsen, Esq.", "title": "Stafford County Commonwealth's Attorney", "address1": "", "address2": ""},
    "Fauquier County": {"name": "Scott Hook, Esq.", "title": "Fauquier County Commonwealth's Attorney", "address1": "29 Ashby Street", "address2": "Warrenton, VA 20186"}
}

# Populate the document
def populate_document(template_path, output_path, replacements):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    doc.save(output_path)
    print(f"Document saved at: {output_path}")

# GUI Setup
root = tk.Tk()
root.title("Expungement Petition Generator")
root.geometry("600x900")

# Create a canvas with a scrollbar
main_frame = tk.Frame(root)
main_frame.pack(fill="both", expand=True)

canvas = tk.Canvas(main_frame)
scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
scrollable_frame = tk.Frame(canvas)

# Configure the scrollbar
scrollable_frame.bind(
    "<Configure>",
    lambda e: canvas.configure(
        scrollregion=canvas.bbox("all")
    )
)

canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
canvas.configure(yscrollcommand=scrollbar.set)

canvas.pack(side="left", fill="both", expand=True)
scrollbar.pack(side="right", fill="y")

def _on_mouse_wheel(event):
    canvas.yview_scroll(-1 * int((event.delta / 120)), "units")

canvas.bind_all("<MouseWheel>", _on_mouse_wheel)

fields = ["{NAME}", "{DOB}", "{Name at Arrest}", "{Arrest Date}", "{Officer Name}", "{Police Department}", "{Charge Name}", "{Code Section}", "{VCC Code}", "{OTN}", "{Court of Final Dispo}", "{Case No}", "{Dispo Date}"]
entries = {}

# Drop-down for County
county_label = tk.Label(scrollable_frame, text="Select County:")
county_label.pack()
county_combo = ttk.Combobox(scrollable_frame, values=list(prosecutor_info.keys()))
county_combo.pack(pady=5)

# Drop-down for Expungement Type
expungement_label = tk.Label(scrollable_frame, text="Select Type of Expungement:")
expungement_label.pack()
expungement_combo = ttk.Combobox(scrollable_frame, values=["Expungement of Right", "Manifest Injustice"])
expungement_combo.pack(pady=5)

# Manifest Injustice Entry
manifest_frame = tk.Frame(scrollable_frame)
manifest_label = tk.Label(manifest_frame, text="Manifest Injustice Details:")
manifest_entry = tk.Entry(manifest_frame, width=50)
manifest_label.pack(side="left")
manifest_entry.pack(side="left", padx=5)
manifest_frame.pack(pady=5, fill="x")

for field in fields:
    frame = tk.Frame(scrollable_frame)
    label = tk.Label(frame, text=field.replace("_", " ").title())
    label.pack(side="left", padx=5, pady=5)
    entry = tk.Entry(frame, width=50)
    entry.pack(side="left", padx=5, pady=5)
    frame.pack(pady=2, fill="x")
    entries[field] = entry

# Drop-down for Final Disposition for the first case
dismissal_label = tk.Label(scrollable_frame, text="Select Final Disposition:")
dismissal_label.pack()
dismissal_combo = ttk.Combobox(scrollable_frame, values=["Nolle Prosequi", "Not Guilty", "Otherwise Dismissed"])
dismissal_combo.pack(pady=5)
entries["{Final Disposition}"] = dismissal_combo

# Add case function
case_count = 1
additional_cases = []

def add_case():
    global case_count
    case_count += 1
    frame = tk.Frame(scrollable_frame)
    case_label = tk.Label(frame, text=f"Additional Case {case_count}")
    case_label.pack(pady=5)

    new_case_fields = ["{Arrest Date}", "{Officer Name}", "{Police Department}", "{Charge Name}", "{Code Section}", "{VCC Code}", "{OTN}"]
    case_entries = {}
    for field in new_case_fields:
        entry_frame = tk.Frame(frame)
        label = tk.Label(entry_frame, text=f"Case {case_count} - " + field.replace("_", " ").title())
        label.pack(side="left", padx=5, pady=5)
        entry = tk.Entry(entry_frame, width=50)
        entry.pack(side="left", padx=5, pady=5)
        entry_frame.pack(pady=2, fill="x")
        case_entries[field] = entry

    # Add Final Disposition dropdown just before Court of Final Dispo
    final_disposition_label = tk.Label(frame, text=f"Case {case_count} - Select Final Disposition:")
    final_disposition_label.pack()
    final_disposition_combo = ttk.Combobox(frame, values=["Nolle Prosequi", "Not Guilty", "Otherwise Dismissed"])
    final_disposition_combo.pack(pady=5)
    case_entries["{Final Disposition}"] = final_disposition_combo

    # Continue with remaining fields after Final Disposition
    remaining_fields = ["{Court of Final Dispo}", "{Case No}", "{Dispo Date}"]
    for field in remaining_fields:
        entry_frame = tk.Frame(frame)
        label = tk.Label(entry_frame, text=f"Case {case_count} - " + field.replace("_", " ").title())
        label.pack(side="left", padx=5, pady=5)
        entry = tk.Entry(entry_frame, width=50)
        entry.pack(side="left", padx=5, pady=5)
        entry_frame.pack(pady=2, fill="x")
        case_entries[field] = entry

    frame.pack(pady=5, fill="x")
    additional_cases.append(case_entries)

# Generate document
def generate_document():
    template_path = filedialog.askopenfilename(title="Select Template", filetypes=[("Word files", "*.docx")])
    output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    county = county_combo.get().strip()
    if not county or county not in prosecutor_info:
        tk.messagebox.showerror("Error", "Please select a valid county from the dropdown.")
        return
    final_disposition = dismissal_combo.get()
    expungement_type = expungement_combo.get()

    # Expungement clause logic
    if expungement_type == "Expungement of Right":
        expungement_clause = ("The Petitioner has no prior criminal record, the aforementioned arrest was a misdemeanor offense, "
                             "and the Commonwealth cannot show good cause to the contrary as to why the petition should not be granted.")
    elif expungement_type == "Manifest Injustice":
        manifest_injustice = manifest_entry.get()
        expungement_clause = (f"The continued existence and possible dissemination of information relating to the charge(s) set forth herein "
                             f"has caused, and may continue to cause, circumstances which constitute a manifest injustice to the Petitioner, "
                             f"and the Commonwealth cannot show good cause to the contrary as to why the petition should not be granted. "
                             f"(to wit: {manifest_injustice}).")
    else:
        expungement_clause = ""

    # Automatically fill prosecutor details based on county
    prosecutor = prosecutor_info[county]["name"]
    prosecutor_title = prosecutor_info[county]["title"]
    prosecutor_address1 = prosecutor_info[county]["address1"]
    prosecutor_address2 = prosecutor_info[county]["address2"]

    replacements = {key: entry.get() for key, entry in entries.items()}
    replacements["{NAME}"] = replacements["{NAME}"].upper()
    replacements["{COUNTY}"] = county.upper()
    replacements["{County2}"] = county.title()
    replacements["{Prosecutor}"] = prosecutor
    replacements["{Prosecutor Title}"] = prosecutor_title
    replacements["{Prosecutor Address 1}"] = prosecutor_address1
    replacements["{Prosecutor Address 2}"] = prosecutor_address2
    replacements["{Final Disposition}"] = final_disposition
    replacements["{Dispo Date}"] = entries["{Dispo Date}"].get()
    replacements["{Type of Expungement}"] = expungement_clause

    # Collect additional cases
    additional_text = ""
    for index, case in enumerate(additional_cases, start=2):
        additional_text += f"\n**Case No. {index}:**\n\n"
        additional_text += f"Date of Arrest: {case['{Arrest Date}'].get()}\n"
        additional_text += f"Arresting Officer: {case['{Officer Name}'].get()}\n"
        additional_text += f"Law Enforcement Agency: {case['{Police Department}'].get()}\n"
        additional_text += f"Charge (Description and Code Section): {case['{Charge Name}'].get()} | {case['{Code Section}'].get()}\n"
        additional_text += f"VCC Code: {case['{VCC Code}'].get()}\n"
        additional_text += f"OTN: {case['{OTN}'].get()}\n"
        additional_text += f"Court of Final Disposition: {case['{Court of Final Dispo}'].get()}\n"
        additional_text += f"Case number: {case['{Case No}'].get()}\n"
        # Safely retrieve Final Disposition and Date
        final_dispo = case['{Final Disposition}'].get() if '{Final Disposition}' in case else ''
        dispo_date = case['{Dispo Date}'].get() if '{Dispo Date}' in case else ''
        additional_text += f"Final Disposition and Date: {final_dispo} on {dispo_date}\n"
        # Include "Certified Copy of Warrant/Summons" only in the petition
        if "Petition" in output_path:
            additional_text += f"Certified Copy of Warrant/Summons attached as Exhibit {index}.\n"

    replacements["{Additional Cases}"] = additional_text

    if template_path and output_path:
        populate_document(template_path, output_path, replacements)

add_case_btn = tk.Button(scrollable_frame, text="Add Another Case", command=add_case)
add_case_btn.pack(pady=10)

submit_btn = tk.Button(scrollable_frame, text="Generate Document", command=generate_document)
submit_btn.pack(pady=20)

root.mainloop()