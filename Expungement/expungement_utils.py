import os
from docx import Document

# Additional imports for PDF parsing and regex
import re
import pdfplumber
from datetime import datetime
import logging
from flask import current_app, jsonify

logger = logging.getLogger(__name__)

# Predefined prosecutor information based on county
prosecutor_info = {
    "Fairfax County": {"name": "Steve Descano, Esq.", "title": "Fairfax County Commonwealth's Attorney", "address1": "4110 Chain Bridge Road", "address2": "Fairfax, VA 22030"},
    "Arlington County": {"name": "Parisa Tafti, Esq.", "title": "Arlington County Commonwealth's Attorney", "address1": "1425 N. Courthouse Rd", "address2": "Arlington, VA 22201"},
    "Prince William County": {"name": "Amy Ashworth, Esq.", "title": "Prince William County Commonwealth's Attorney", "address1": "9311 Lee Ave", "address2": "Manassas, VA 20110"},
    "Loudoun County": {"name": "Robert Anderson, Esq.", "title": "Loudoun County Commonwealth's Attorney", "address1": "20 E Market St", "address2": "Leesburg, VA 20176"},
    "City of Alexandria": {"name": "Bryan Porter, Esq.", "title": "Alexandria City Commonwealth's Attorney", "address1": "520 King Street", "address2": "Alexandria, VA 22314"},
    "Stafford County": {"name": "Eric Olsen, Esq.", "title": "Stafford County Commonwealth's Attorney", "address1": "", "address2": ""},
    "Fauquier County": {"name": "Scott Hook, Esq.", "title": "Fauquier County Commonwealth's Attorney", "address1": "29 Ashby Street", "address2": "Warrenton, VA 20186"}
}

def populate_document(template_path, output_path, replacements):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    doc.save(output_path)

def parse_officer_name(raw_name):
    val = raw_name.strip()
    parts = val.split(",", 1)
    if len(parts) == 2:
        last, first = parts
        val = f"{last.strip().title()}, {first.strip().title()}"
    else:
        val = val.title()
    return val.replace("&", "&amp;")

def safe_search(pattern, text, flags=0):
    try:
        return re.search(pattern, text, flags)
    except re.error as e:
        logger.warning(f"Regex error for pattern '{pattern}': {e}")
        return None

def extract_expungement_data(filepath, case_index=None):
    text = ""
    logger.debug(f"Extracting data from PDF: {filepath}")

    # Assumes filepath is directly accessible (absolute or temp file path from upload)
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
                logger.debug(f"Extracted text from page: {page.page_number}")

    field_patterns = [
        ("name", [
            r"Defendant:\s*([^\n\r]+)"
        ]),
        ("dob", [
            r"DOB:\s*(\d{2}/\d{2}/\d{4})",
            r"DOB:\s*(\d{2}/\d{2}/\*{4})"
        ]),
        ("officer_name", [
            r"Complainant\s*:\s*([\w\s.,'-]+)",
            r"Complainant:\s*([^\n\r]+)"
        ]),
        ("arrest_date", [
            r"Arrest Date\s*:\s*(\d{2}/\d{2}/\d{4})",
            r"Arrest Date:\s*(\d{2}/\d{2}/\d{4})"
        ]),
        ("dispo_date", [
            r"Disposition Date\s*:\s*(\d{2}/\d{2}/\d{4})"
        ]),
        ("charge_name", [
            r"Charge\s*:\s*([^\n\r]+)"
        ]),
        ("code_section", [
            r"Code\s*Section\s*:\s*([\d\.]+-\d+)",
            r"Code Section:\s*(.+)"
        ]),
        ("otn", [
            r"OffenseTracking/Processing#\s*:\s*(\S+)",
            r"Offense Tracking/Process #:\s*(\S+)"
        ]),
        ("case_no", [
            r"Case\s*(?:No|#|Number)\s*[:\-]?\s*([A-Z0-9\-]+)",
            r"Case\s*#\s*:\s*(\S+)",
            r"Case #:\s*(GC\d{8}-\d{2})",
        ]),
        ("final_dispo", [
            r"Disposition:\s*([A-Z\s]+)",
            r"Disposition:\s*(NOLLE PROSEQUI|DISMISSED|GUILTY|NOT GUILTY)"
        ]),
        ("court_dispo", [
            r"^([A-Z][a-z]+ (County|City) (Juvenile and Domestic Relations District Court|General District Court|Circuit Court))",
            r"Return to Search Results\s*([\w\s]+(?:General District Court|Circuit Court|Juvenile and Domestic Relations District Court))",
            r"(\w+\s+(?:County|City)\s+(?:General District Court|Circuit Court|Juvenile and Domestic Relations District Court))",
            r"((?:General District Court|Circuit Court|Juvenile and Domestic Relations District Court)\s*[-–]?\s*\w+)",
            r"(\w+\s+Juvenile and Domestic Relations District Court)",
            r"(Fairfax County Juvenile and Domestic Relations District Court)"
        ]),
        ("police_department", [
            r"Arresting Agency\s*:\s*([^\n\r]+)",
            r"Police Department\s*:\s*([^\n\r]+)"
        ]),
    ]

    result = {}

    def assign_value(key, val):
        val = val.replace("&", "&amp;")
        if case_index is not None:
            result[f"case_{case_index}_{key}"] = val
        else:
            result[key] = val

    for key, patterns in field_patterns:
        match = None
        val = ""
        for pattern in patterns:
            match = safe_search(pattern, text, re.MULTILINE)
            if match:
                val = match.group(1).strip()
                break

        if key == "arrest_date":
            # Will handle fallback later
            if val:
                assign_value(key, val)
            else:
                assign_value(key, "")
            continue

        if match:
            if key == "name":
                if case_index is not None and case_index != 1:
                    continue
                # Handle names in format: Last, First Middle
                if "," in val:
                    last, rest = val.split(",", 1)
                    name_parts = rest.strip().split()
                    formatted_name = " ".join(name_parts + [last.strip()])
                else:
                    formatted_name = val.strip()
                if formatted_name:
                    formatted_name = " ".join(word.capitalize() for word in formatted_name.split())
                assign_value(key, formatted_name)
                assign_value("name_arrest", formatted_name)
            elif key == "dob":
                if case_index is not None and case_index != 1:
                    continue
                try:
                    val = datetime.strptime(val, "%m/%d/%Y").strftime("%Y-%m-%d")
                except ValueError:
                    pass
                assign_value(key, val)
            elif key == "charge_name":
                val = re.split(r"Offense Tracking|Process|Code\s*Section|Case\s*Type", val)[0].strip()
                val = " ".join(word.capitalize() for word in val.split())
                assign_value(key, val)
            elif key == "court_dispo":
                val = val.strip().replace('\n', ' ').strip()
                assign_value(key, val)
            elif key == "officer_name":
                val = re.split(r"Amended|Hearing|Disposition", val)[0].strip()
                val = parse_officer_name(val)
                assign_value(key, val)
            elif key == "otn":
                val = re.sub(r"Summons.*", "", val).strip()
                val = re.split(r"Case", val)[0].strip()
                assign_value(key, val)
            elif key == "final_dispo":
                val = re.split(r"Sentence", val)[0].strip()
                val = " ".join(word.capitalize() for word in val.split())
                # Remove trailing solitary 'S' or 'S.' if it's an unintended character
                if val.strip().upper() in ["", "."]:
                    assign_value(key, "")
                else:
                    # Only remove trailing S if "Nolle Pro" is in value, to handle "Nolle ProS"
                    val = val.rstrip("S").strip() if val.upper().endswith("S") and "Nolle Pro" in val.title() else val
                    assign_value(key, val)
            elif key == "code_section":
                val = re.split(r"Charge", val)[0].strip()
                if not val.startswith("Va. Code"):
                    val = f"Va. Code § {val}"
                assign_value(key, val)
            elif key == "dispo_date":
                assign_value(key, val)
            elif key == "police_department":
                assign_value(key, val)
            else:
                assign_value(key, val)
        else:
            if key == "officer_name":
                fallback = safe_search(r"Complainant\s*[:\-]?\s*([A-Z][a-z]+,\s*[A-Z]\s*[A-Z]?)", text)
                fallback_val = ""
                if fallback:
                    fallback_val = fallback.group(1).strip()
                val = parse_officer_name(fallback_val) if fallback_val else ""
                assign_value(key, val)
            elif key != "arrest_date":  # already handled separately
                if case_index is not None:
                    if case_index != 1 and key in ["name", "name_arrest", "dob"]:
                        continue
                    else:
                        assign_value(key, "")
                else:
                    assign_value(key, "")

    # Fallback logic for arrest_date
    if (case_index is not None and f"case_{case_index}_arrest_date" not in result) or (case_index is None and "arrest_date" not in result):
        arrest_fallback = safe_search(r"Arrest(?:ed)? Date\s*[:\-]?\s*(\d{2}/\d{2}/\d{4})", text, re.IGNORECASE)
        fallback_val = arrest_fallback.group(1).strip() if arrest_fallback else ""
        assign_key = f"case_{case_index}_arrest_date" if case_index is not None else "arrest_date"
        result[assign_key] = fallback_val.replace("&", "&amp;") if fallback_val else ""

    # Fallback for court_dispo if not matched and case_index is not None or no value found
    court_key = f"case_{case_index}_court_dispo" if case_index is not None else "court_dispo"
    if court_key not in result or not result[court_key]:
        fallback_court = safe_search(r"([A-Z][a-z]+ (County|City) Juvenile and Domestic Relations District Court)", text)
        if fallback_court:
            result[court_key] = fallback_court.group(1).strip().replace("&", "&amp;")
        elif case_index is not None:
            # Final fallback: extract court name preceding "(details)" from top of first page
            try:
                with pdfplumber.open(filepath) as pdf:
                    first_page_text = pdf.pages[0].extract_text()
                    if first_page_text:
                        for line in first_page_text.split("\n"):
                            if "(details" in line.lower():
                                court_name = line.split("(details")[0].strip()
                                if "court" in court_name.lower():
                                    result[court_key] = court_name.replace("&", "&amp;")
                                break
            except Exception as e:
                logger.exception(f"Fallback court_dispo scan failed: {e}")

    # Fallback logic for dispo_date if empty: collect all hearing dates in Hearing Information section
    dispo_key = f"case_{case_index}_dispo_date" if case_index is not None else "dispo_date"
    if dispo_key not in result or not result[dispo_key]:
        hearing_section = safe_search(r"Hearing Information(.*?)Disposition Information", text, re.DOTALL | re.IGNORECASE)
        if hearing_section:
            dates = re.findall(r"\b(\d{2}/\d{2}/\d{4})\b", hearing_section.group(1))
            if dates:
                first_date = dates[0]
                result[dispo_key] = first_date.replace("&", "&amp;")

    # If case_index is set, filter keys to only those with the case_{case_index}_ prefix
    if case_index is not None:
        filtered_result = {}
        for k, v in result.items():
            if k.startswith(f"case_{case_index}_"):
                filtered_result[k] = v
        logger.debug(f"Extraction result: {filtered_result}")
        return filtered_result

    if case_index is None:
        for required_key in ["name", "name_arrest", "dob", "final_dispo"]:
            if required_key not in result:
                result[required_key] = ""

    logger.debug(f"Extraction result: {result}")
    return result

def map_case_keys(extracted, case_index=1):
    mapping = {
        f"case_{case_index}_name": "full_legal_name",
        f"case_{case_index}_name_arrest": "name_arrest",
        f"case_{case_index}_dob": "dob",
        f"case_{case_index}_officer_name": "officer_name",
        f"case_{case_index}_arrest_date": "arrest_date",
        f"case_{case_index}_dispo_date": "dispo_date",
        f"case_{case_index}_charge_name": "charge_name",
        f"case_{case_index}_code_section": "code_section",
        f"case_{case_index}_otn": "otn",
        f"case_{case_index}_case_no": "case_no",
        f"case_{case_index}_final_dispo": "final_dispo",
        f"case_{case_index}_court_dispo": "court_dispo",
        f"case_{case_index}_police_department": "police_department",
    }
    return {mapping.get(k, k): v for k, v in extracted.items() if k in mapping}

if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) < 2:
        print("Usage: python expungement_utils.py <PDF_FILE_1> <PDF_FILE_2> ...")
        sys.exit(1)

    pdf_paths = sys.argv[1:]
    all_results = []

    for idx, path in enumerate(pdf_paths, start=1):
        case_data = extract_expungement_data(path, case_index=idx)
        all_results.append(case_data)

    print(json.dumps(all_results, indent=2))


# Function to register after_request handler for JSON Content-Type

def register_after_request(app):
    @app.after_request
    def ensure_json_content_type(response):
        if response.is_json or response.mimetype == "application/json":
            response.headers["Content-Type"] = "application/json"
        return response


def extract_multiple_cases_data(filepaths):
    all_data = []
    for idx, path in enumerate(filepaths, start=1):
        case_data = extract_expungement_data(path, case_index=idx)
        all_data.append(case_data)
    return all_data

